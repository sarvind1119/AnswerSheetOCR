from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from .analytics import build_analytics
from .config import REPORTS_DIR
from .models import RunRecord


def generate_docx_report(record: RunRecord, output_path: Path | None = None) -> Path:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    target = output_path or REPORTS_DIR / f"{record.metadata.run_id}.docx"

    analytics = build_analytics(record)
    document = Document()
    document.add_heading("Handwritten Answer Sheet Digitization Report", level=1)

    document.add_heading("Run Details", level=2)
    details = document.add_table(rows=0, cols=2)
    for key, value in [
        ("Source PDF", record.metadata.source_pdf_name),
        ("Run ID", record.metadata.run_id),
        ("Model", record.metadata.model),
        ("DPI", str(record.metadata.dpi)),
        ("Generated", datetime.now().strftime("%Y-%m-%d %H:%M")),
    ]:
        row = details.add_row().cells
        row[0].text = key
        row[1].text = value

    document.add_heading("Analytics Summary", level=2)
    summary = document.add_table(rows=0, cols=2)
    for key, value in [
        ("Rendered pages", analytics["page_count"]),
        ("OCR processed pages", analytics["ocr_page_count"]),
        ("Question blocks", analytics["question_count"]),
        ("Uncertainty flags", analytics["uncertainty_count"]),
        ("Missing question numbers", analytics["missing_question_numbers"]),
        ("Detected languages", ", ".join(analytics["detected_languages"].keys()) or "None"),
    ]:
        row = summary.add_row().cells
        row[0].text = str(key)
        row[1].text = str(value)

    document.add_heading("Question-wise Text", level=2)
    for page in sorted(record.pages, key=lambda item: item.page_number):
        document.add_heading(f"Page {page.page_number}", level=3)
        if page.detected_languages:
            document.add_paragraph("Languages: " + ", ".join(page.detected_languages))

        if not page.questions:
            document.add_paragraph("No question blocks extracted.")
            continue

        for index, question in enumerate(page.questions, start=1):
            label = question.question_number or f"Unnumbered block {index}"
            document.add_heading(f"Question {label}", level=4)
            paragraph = document.add_paragraph()
            paragraph.add_run(question.review_text or "[No text extracted]")

            if question.corrected_text.strip():
                document.add_paragraph("Reviewed correction applied.")
            if question.structure_elements:
                document.add_paragraph(
                    "Structure: " + "; ".join(question.structure_elements)
                )
            if question.margin_notes:
                document.add_paragraph("Margin notes: " + "; ".join(question.margin_notes))
            if question.uncertainty_flags:
                document.add_paragraph(
                    "Uncertainty flags: " + "; ".join(question.uncertainty_flags)
                )
            if question.page_refs:
                document.add_paragraph(
                    "Page references: " + ", ".join(map(str, question.page_refs))
                )

    document.save(target)
    return target
